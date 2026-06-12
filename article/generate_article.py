"""Generator draft artikel ilmiah (.docx) untuk UAS INF322.

Membuat dokumen berformat Cambria 11 (bodi) / Cambria 10 (judul gambar-tabel),
lengkap dengan tabel kebutuhan, tabel pengujian kotak hitam, narasi tiap fitur,
dan SLOT GAMBAR berbingkai yang tinggal ditempeli screenshot.

Jalankan:  python generate_article.py
Hasil   :  Artikel_INF322_VisionPrep.docx
"""

import os

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BODY_FONT = "Cambria"
BODY_SIZE = 11
CAPTION_SIZE = 10

fig_no = 0
tab_no = 0


def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def style_document(doc):
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_SIZE)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    for name, size in (("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)):
        st = doc.styles[name]
        st.font.name = BODY_FONT
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)


def add_body(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph(text)
    p.alignment = align
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(CAPTION_SIZE)
    run.italic = True
    return p


def add_image_slot(doc, caption, hint="Tempel gambar/screenshot di sini",
                   image=None, width_in=6.0):
    """Sisipkan gambar jika tersedia; jika tidak, tampilkan kotak slot kosong.

    Selalu diikuti judul gambar (Cambria 10).
    """
    global fig_no
    fig_no += 1
    if image and os.path.exists(image):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(image, width=Inches(width_in))
    else:
        table = doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell = table.cell(0, 0)
        set_cell_shading(cell, "F2F2F2")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[ {hint} ]")
        run.font.name = BODY_FONT
        run.font.size = Pt(CAPTION_SIZE)
        run.font.color.rgb = RGBColor(0x90, 0x90, 0x90)
        for _ in range(4):
            cell.add_paragraph()
    add_caption(doc, f"Gambar {fig_no}. {caption}")


def add_table_caption(doc, caption):
    global tab_no
    tab_no += 1
    add_caption(doc, f"Tabel {tab_no}. {caption}")


def fill_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].paragraphs[0].add_run(h).bold = True
        set_cell_shading(hdr[i], "D9E2F3")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return table


def build():
    doc = Document()
    style_document(doc)

    # ---------------------------------------------------------------- JUDUL
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Pengembangan Tool Preprocessing Gambar Semi-Otomatis "
                        "untuk Pembuatan Dataset Model Vision Berbasis ANN")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = BODY_FONT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        "Azi Dwipurnama (2023071025), Benediktus Aditya Pratama (2023071027), "
        "Marcelinus Nugraha Pratama (2023071028),\n"
        "Muhamad Al Fariz (2023071029), Dava Pangestu Putera (2023071030), "
        "Muhamad Aulia Rifki (2023071026)\n"
        "Program Studi S1 Informatika — INF322 Image Processing\n"
        "Universitas Pembangunan Jaya")

    # ----------------------------------------------------------- PENDAHULUAN
    doc.add_heading("1. Pendahuluan", level=1)
    add_body(doc,
        "Kebutuhan pada penelitian ini berasal dari sebuah perusahaan yang "
        "bergerak di bidang visi komputer (computer vision). Perusahaan "
        "tersebut memerlukan sebuah tool untuk melakukan preprocessing "
        "sejumlah gambar guna menyiapkan dataset bagi suatu model vision. "
        "Tool yang dibutuhkan harus bekerja secara semi-otomatis, yaitu "
        "pengguna cukup mengunggah gambar berformat .jpg ke sebuah folder "
        "kemudian menjalankan lima tahap pengolahan: color-to-grayscale "
        "conversion, resizing, binarization, creating dataset (termasuk "
        "label), dan randomize dataset. Karena alasan komersialisasi, "
        "aplikasi dibangun tanpa pustaka pihak ketiga kecuali NumPy dan "
        "Matplotlib dasar, serta berjalan pada desktop Windows 11. Penelitian "
        "ini bertujuan membangun tool tersebut dan membuktikan bahwa dataset "
        "yang dihasilkannya valid dengan menggunakannya untuk melatih sebuah "
        "model Artificial Neural Network (ANN).")

    # ----------------------------------------------------------------- METODE
    doc.add_heading("2. Metode", level=1)
    add_body(doc, "Penelitian ini dikerjakan dengan melibatkan metode-metode "
                  "berikut ini.")

    doc.add_heading("2.1 Alur Kerja", level=2)
    add_body(doc,
        "Alur kerja penelitian dimulai dari pengambilan informasi kebutuhan "
        "pengguna, dilanjutkan dengan perancangan dan pembangunan program "
        "untuk setiap fitur yang diminta, lalu diakhiri dengan pengujian "
        "aplikasi menggunakan metode kotak hitam (black-box testing). "
        "Keseluruhan tahapan tersebut digambarkan pada diagram alur berikut.")
    add_image_slot(doc, "Diagram alur kerja penelitian.",
                   image="diagram_alur.png", width_in=6.4)

    doc.add_heading("2.2 Pemahaman Kebutuhan Pengguna", level=2)
    add_body(doc,
        "Kebutuhan pengguna dirangkum dari hasil pertemuan dengan tim "
        "perusahaan. Ringkasan kebutuhan tersebut disajikan pada Tabel 1.")
    fill_table(doc,
        ["No", "Kebutuhan Sistem", "Spesifikasi Teknis", "Pendekatan Solusi"],
        [
            ["1",
             "Sistem membutuhkan preprocessing semi-otomatis tahap demi tahap.",
             "Pengguna memilih tahap lalu menekan tombol Start untuk "
             "menjalankannya.",
             "Membangun GUI dengan tombol untuk tiap tahap beserta penanganan "
             "event-nya."],
            ["2",
             "Sistem membutuhkan konversi gambar berwarna menjadi grayscale.",
             "Program mengubah gambar RGB menjadi satu kanal grayscale.",
             "Menggunakan rumus luminance terbobot 0,299R + 0,587G + 0,114B."],
            ["3",
             "Sistem membutuhkan perkecilan ukuran gambar (resizing).",
             "Pengguna memilih metode pooling dan mengisi ukuran (Row, Col).",
             "Menyediakan Average Pooling dan Max Pooling dengan ukuran hasil "
             "yang dapat diisi pengguna."],
            ["4",
             "Sistem membutuhkan binarisasi gambar.",
             "Program menghasilkan gambar hitam-putih murni.",
             "Menggunakan inversi citra dilanjutkan thresholding."],
            ["5",
             "Sistem membutuhkan pembuatan dataset beserta label.",
             "Program menyusun matriks piksel dan label tiap sampel.",
             "Melakukan flatten piksel dan one-hot encoding; kelas dideteksi "
             "dari nama berkas."],
            ["6",
             "Sistem membutuhkan pengacakan dataset.",
             "Program mengacak urutan sampel secara konsisten.",
             "Mengacak deret indeks lalu menatanya pada inputs dan labels "
             "sekaligus."],
            ["7",
             "Sistem membutuhkan pembuktian validitas dataset.",
             "Dataset dapat digunakan untuk melatih model ANN.",
             "Membangun model ANN satu hidden layer berbasis NumPy."],
            ["8",
             "Sistem dibangun tanpa pustaka pihak ketiga.",
             "Hanya boleh memakai NumPy dan Matplotlib dasar.",
             "Mengimplementasikan algoritma secara manual dan memakai Tkinter "
             "bawaan Python."],
            ["9",
             "Sistem berjalan sebagai aplikasi desktop Windows 11.",
             "Program kompatibel dengan Windows 11.",
             "Menggunakan Python dan Tkinter yang tidak bergantung pada OS "
             "tertentu."],
            ["10",
             "Sistem menampilkan masukan dan keluaran tiap tahap.",
             "Hasil tiap tahap berada pada folder yang sama dengan masukan.",
             "Menyimpan keluaran dengan akhiran _gs, _rs, _bin, serta berkas "
             ".npy pada folder tersebut."],
        ])
    add_table_caption(doc, "Ringkasan Kebutuhan Pengguna.")

    doc.add_heading("2.3 Color-to-Grayscale Conversion", level=2)
    add_body(doc,
        "Konversi citra berwarna menjadi grayscale dilakukan dengan rumus "
        "luminance terbobot, yaitu gray = 0,299R + 0,587G + 0,114B (standar "
        "ITU-R BT.601). Pembobotan ini lebih sesuai dengan sensitivitas mata "
        "manusia dibandingkan rata-rata sederhana, sehingga kontras objek "
        "terhadap latar lebih terjaga. Nilai grayscale dihitung untuk setiap "
        "piksel dan disimpan sebagai citra satu kanal.")

    doc.add_heading("2.4 Resizing", level=2)
    add_body(doc,
        "Resizing memperkecil ukuran citra melalui operasi pooling. Citra "
        "dibagi menjadi wilayah-wilayah (pooling region) yang ukurannya "
        "proporsional terhadap rasio ukuran asal dan ukuran hasil. Pada "
        "Average Pooling, nilai setiap piksel hasil diperoleh dari rata-rata "
        "seluruh piksel pada wilayahnya, sehingga citra cenderung lebih halus. "
        "Pada Max Pooling, nilai piksel hasil diambil dari nilai maksimum pada "
        "wilayahnya, sehingga bagian paling terang (objek) lebih dipertahankan. "
        "Pengguna dapat menentukan sendiri ukuran hasil melalui isian Row dan "
        "Col, misalnya 20 x 30.")
    add_image_slot(doc, "Ilustrasi Average Pooling dan Max Pooling.",
                   image="ilustrasi_pooling.png", width_in=6.2)

    doc.add_heading("2.5 Binarization", level=2)
    add_body(doc,
        "Binarization mengubah citra grayscale menjadi citra hitam-putih "
        "murni. Citra terlebih dahulu diinversi (255 - nilai) agar latar "
        "menjadi hitam dan objek menjadi putih, kemudian di-threshold: piksel "
        "bernilai lebih besar atau sama dengan ambang (default 100) diset 255 "
        "(putih), selebihnya diset 0 (hitam). Hasilnya berupa citra biner yang "
        "menonjolkan bentuk objek.")

    doc.add_heading("2.6 Creating Dataset", level=2)
    add_body(doc,
        "Setiap citra biner diubah menjadi satu baris dataset dengan cara "
        "flatten (penderetan piksel secara baris demi baris). Kolom ke-0 "
        "menyimpan nomor sampel, sedangkan kolom berikutnya menyimpan nilai "
        "piksel. Label dibuat dalam bentuk one-hot encoding, dengan kelas "
        "dideteksi otomatis dari awalan nama berkas (misalnya berkas "
        "“katakana_a (3)” termasuk kelas “katakana_a”). "
        "Dataset masukan dan label disimpan dalam berkas .npy.")

    doc.add_heading("2.7 Randomize Dataset", level=2)
    add_body(doc,
        "Pengacakan dilakukan dengan membangkitkan deret indeks 0..n-1 lalu "
        "mengacaknya. Deret indeks yang sama digunakan untuk menata ulang "
        "baris dataset masukan dan label secara bersamaan, sehingga pasangan "
        "gambar dan labelnya tetap konsisten. Pengacakan penting agar proses "
        "pelatihan ANN tidak bias terhadap urutan kelas.")

    doc.add_heading("2.8 Penyiapan Graphical User Interface (GUI)", level=2)
    add_body(doc,
        "Antarmuka dibangun menggunakan pustaka Tkinter bawaan Python sehingga "
        "tidak melanggar batasan pustaka. GUI menyediakan tombol pemilih "
        "folder, daftar tahap yang dapat dipilih, dan tombol Start untuk "
        "menjalankan tahap terpilih, sesuai pola interaksi yang diminta "
        "pengguna. Antarmuka dirancang bertema gelap dengan satu warna aksen "
        "agar nyaman dipandang. Untuk memudahkan verifikasi, GUI dilengkapi "
        "panel preview yang menampilkan citra INPUT dan OUTPUT dari tahap yang "
        "sedang dipilih, serta sebuah konsol log yang menampilkan jalannya "
        "proses. Proses berat dijalankan pada thread terpisah agar antarmuka "
        "tetap responsif. Hasil tiap tahap disimpan pada folder yang sama "
        "dengan gambar masukan.")

    # --------------------------------------------------- HASIL & PEMBAHASAN
    doc.add_heading("3. Hasil dan Pembahasan", level=1)
    add_body(doc, "Hasil-hasil dari pelaksanaan project ini disampaikan "
                  "sebagai berikut ini.")

    doc.add_heading("3.1 Hasil", level=2)

    sections = [
        ("Color-to-Grayscale Conversion",
         "Fitur grayscale berhasil mengubah seluruh citra berwarna menjadi "
         "citra grayscale."),
        ("Resizing",
         "Fitur resizing berhasil memperkecil citra ke ukuran yang ditentukan "
         "pengguna, baik dengan Average Pooling maupun Max Pooling."),
        ("Binarization",
         "Fitur binarization berhasil menghasilkan citra biner dengan objek "
         "berwarna putih pada latar hitam."),
        ("Creating Dataset",
         "Fitur creating dataset berhasil membentuk berkas dataset masukan dan "
         "label dengan kelas terdeteksi otomatis."),
        ("Randomize Dataset",
         "Fitur randomize berhasil mengacak urutan sampel dataset secara "
         "konsisten antara masukan dan label."),
        ("Training & Testing ANN",
         "Dataset yang dihasilkan berhasil digunakan untuk melatih model ANN "
         "hingga mencapai akurasi yang tinggi, dan model dapat mengklasifikasi "
         "data uji baik insample maupun gambar custom di luar dataset."),
    ]
    # Figur masukan-keluaran hasil generate untuk tiap fitur (file, lebar inci).
    figs = {
        "Color-to-Grayscale Conversion": ("io_grayscale.png", 5.2),
        "Resizing": ("io_resizing.png", 5.2),
        "Binarization": ("io_binarization.png", 5.2),
        "Creating Dataset": ("montase_dataset.png", 6.0),
        "Randomize Dataset": ("randomize_order.png", 6.2),
        "Training & Testing ANN": ("grafik_akurasi.png", 5.2),
    }
    for judul, kalimat in sections:
        doc.add_heading(judul, level=3)
        add_body(doc, kalimat)
        add_image_slot(doc, f"Tangkapan layar kode program fitur {judul}.",
                       "Tempel screenshot kode di sini")
        fig = figs.get(judul)
        if fig:
            add_image_slot(doc, f"Masukan dan keluaran fitur {judul}.",
                           image=fig[0], width_in=fig[1])
        else:
            add_image_slot(doc, f"Contoh masukan dan keluaran fitur {judul}.",
                           "Tempel gambar input/output di sini")

    doc.add_heading("Graphical User Interface", level=3)
    add_body(doc, "Tampilan antarmuka aplikasi disajikan pada gambar berikut.")
    add_image_slot(doc, "Tampilan utama GUI aplikasi.",
                   "Tempel tangkapan layar GUI di sini")

    doc.add_heading("Hasil Pengujian Kotak Hitam", level=3)
    add_body(doc, "Pengujian kotak hitam dilakukan terhadap setiap fitur. "
                  "Hasilnya disajikan pada Tabel 2.")
    fill_table(doc,
        ["No", "Skenario Uji", "Masukan", "Keluaran Diharapkan",
         "Keluaran Aktual", "Status"],
        [
            ["1", "Color-to-Grayscale", "Folder berisi gambar .jpg berwarna",
             "Tiap gambar menjadi grayscale (_gs.jpg)", "Sesuai", "Lulus"],
            ["2", "Resizing Average Pooling", "Gambar _gs.jpg, Row=20 Col=30",
             "Gambar berukuran 20x30 (_rs.jpg)", "Sesuai", "Lulus"],
            ["3", "Resizing Max Pooling", "Gambar _gs.jpg, Row=20 Col=30",
             "Gambar berukuran 20x30 (_rs.jpg)", "Sesuai", "Lulus"],
            ["4", "Binarization", "Gambar _rs.jpg, threshold=70",
             "Gambar biner hitam-putih (_bin.jpg)", "Sesuai", "Lulus"],
            ["5", "Creating Dataset", "Gambar _bin.jpg beberapa kelas",
             "Berkas inputs_*.npy & labels_*.npy", "Sesuai", "Lulus"],
            ["6", "Randomize Dataset", "Berkas dataset .npy",
             "Berkas random_*.npy dengan urutan teracak", "Sesuai", "Lulus"],
            ["7", "Train ANN", "Dataset acak (60 sampel, 150 epoch)",
             "Akurasi pelatihan meningkat tinggi", "Akurasi 100%", "Lulus"],
            ["8", "Test ANN insample", "Nomor sampel dari dataset",
             "Prediksi sesuai label", "Sesuai", "Lulus"],
            ["9", "Test ANN custom", "Gambar .jpg luar dataset",
             "Prediksi kelas yang masuk akal", "Sesuai", "Lulus"],
            ["10", "Validasi masukan kosong", "Folder belum dipilih",
             "Pesan peringatan pada konsol", "Sesuai", "Lulus"],
        ])
    add_table_caption(doc, "Hasil Pengujian Kotak Hitam.")

    doc.add_heading("3.2 Pembahasan", level=2)
    add_body(doc,
        "Berdasarkan hasil di atas, kelima fitur preprocessing yang diminta "
        "pengguna berhasil dibangun dan bekerja sesuai harapan. Pengujian "
        "dilakukan menggunakan dataset huruf katakana tulisan tangan yang "
        "terdiri dari tiga kelas, yaitu A, I, dan U, masing-masing sebanyak "
        "20 sampel, sehingga totalnya 60 gambar. Karena goresan tulisan relatif "
        "tipis, nilai ambang binarisasi diatur ke 70 (lebih rendah dari nilai "
        "bawaan 100) agar bentuk huruf tetap terjaga setelah citra diperkecil "
        "ke ukuran 20 x 30 dengan Average Pooling. Walaupun sebagian citra "
        "biner mengandung sedikit derau, model ANN tetap mampu membedakan "
        "ketiga kelas karena bentuk ketiga huruf cukup berbeda. Validitas "
        "dataset terbukti dari keberhasilan model ANN mencapai akurasi 100% "
        "pada data latih setelah 150 epoch, serta kemampuannya mengklasifikasi "
        "gambar di luar dataset (outsample). Seluruh proses dibangun hanya "
        "dengan NumPy dan Matplotlib dasar serta Tkinter bawaan, sehingga "
        "memenuhi batasan komersialisasi yang ditetapkan.")

    # ------------------------------------------------------------- KESIMPULAN
    doc.add_heading("4. Kesimpulan", level=1)
    add_body(doc,
        "Penelitian ini berhasil membangun sebuah tool preprocessing gambar "
        "semi-otomatis yang mencakup lima tahap, yaitu color-to-grayscale "
        "conversion, resizing dengan average/max pooling, binarization, "
        "creating dataset beserta label, dan randomize dataset, seluruhnya "
        "dibangun hanya dengan NumPy dan Matplotlib dasar serta antarmuka "
        "Tkinter. Dataset yang dihasilkan terbukti valid karena berhasil "
        "digunakan untuk melatih model ANN hingga mencapai akurasi tinggi. "
        "Dengan demikian, kelima fitur bekerja dengan baik dan tool yang "
        "dibangun telah bekerja sebagaimana diharapkan.")
    link = doc.add_paragraph()
    link.add_run("URL Video: ").bold = True
    link.add_run("[Tempel tautan video GDrive Anda di sini]")

    doc.save("Artikel_INF322_VisionPrep.docx")
    print("Tersimpan: Artikel_INF322_VisionPrep.docx")


if __name__ == "__main__":
    build()
